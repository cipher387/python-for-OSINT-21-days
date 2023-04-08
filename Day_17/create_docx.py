from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Report', 0)

document.add_heading('Report', level=1)


document.add_paragraph(
    'Some text in report'
)

document.add_page_break()

document.add_picture('histogram.png', width=Inches(1.25))

document.save('report.docx')
